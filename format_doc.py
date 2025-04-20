import os
from docx import Document
from docx.shared import Pt

def save_text_to_docx(text, query, output_path=None):
    
    # Ensure we create a new document
    doc = Document()

    # Add a title
    title = doc.add_paragraph()
    title_run = title.add_run("Extracted Information")
    title_run.bold = True
    title_run.font.size = Pt(16)

    # Add the query
    doc.add_paragraph(f"Query: {query}", style="Intense Quote")

 
    for paragraph in text.split("\n\n"):
        doc.add_paragraph(paragraph)

    
    if output_path is None:
        output_path = os.path.abspath("formatted_response.docx")

    try:
        doc.save(output_path)
        print(f"✅ Document saved successfully at {output_path}")
    except Exception as e:
        print(f"🚨 Error saving document: {e}")

    return output_path


